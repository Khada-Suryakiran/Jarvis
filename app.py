import os
from flask import Flask, request, jsonify, render_template, send_file
import google.generativeai as genai
from gtts import gTTS
from io import BytesIO
from fpdf import FPDF
from docx import Document
from dotenv import load_dotenv
import openai
from datetime import datetime

load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_GENAI_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not GOOGLE_API_KEY or not OPENAI_API_KEY:
    raise ValueError("API keys are not set in environment variables")

genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')
openai.api_key = OPENAI_API_KEY


tmpl_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
app = Flask(__name__, template_folder=tmpl_dir, static_folder="static")

chat_history = []

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    try:
        text = request.json['text']
        response = model.generate_content(text)
        generated_text = response.text
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        chat_history.append({'prompt': text, 'response': generated_text, 'timestamp': timestamp})
        return jsonify({'generated_text': generated_text, 'timestamp': timestamp})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate-audio', methods=['POST'])
def generate_audio():
    try:
        text = request.json['text']
        tts = gTTS(text)
        audio_data = BytesIO()
        tts.write_to_fp(audio_data)
        audio_data.seek(0)
        return send_file(audio_data, mimetype='audio/mpeg', as_attachment=False)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download-pdf')
def download_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for entry in chat_history:
        pdf.multi_cell(0, 10, f"User: {entry['prompt']}\nJarvis: {entry['response']}\nTime: {entry['timestamp']}\n")
    pdf_path = "chat_history.pdf"
    pdf.output(pdf_path)
    return send_file(pdf_path, as_attachment=True)

@app.route('/download-docx')
def download_docx():
    doc = Document()
    doc.add_heading('Chat History', 0)
    for entry in chat_history:
        doc.add_paragraph(f"User: {entry['prompt']}")
        doc.add_paragraph(f"Jarvis: {entry['response']}")
        doc.add_paragraph(f"Time: {entry['timestamp']}")
        doc.add_paragraph("")
    doc_path = "chat_history.docx"
    doc.save(doc_path)
    return send_file(doc_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
